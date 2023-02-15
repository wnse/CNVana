# ---
# jupyter:
#   jupytext:
#     formats: ipynb,py:percent
#     text_representation:
#       extension: .py
#       format_name: percent
#       format_version: '1.3'
#       jupytext_version: 1.13.7
#   kernelspec:
#     display_name: Python 3 (ipykernel)
#     language: python
#     name: python3
# ---

# %%
import os
import re
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
import pandas as pd
import logging
import argparse


# %%
def fill_cell(cell, content, font_name='微软雅黑', font_size=10, color=RGBColor(0, 0, 0), bold=False, text_alignment=None):
    run = cell.paragraphs[0].add_run(str(content))
    if text_alignment:
        cell.paragraphs[0].alignment = text_alignment
    cell.paragraphs[0].paragraph_format.line_spacing=Pt(18)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color

def add_row(table, row_no, height=None):
    while len(table.rows) < row_no+1:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT.AT_LEAST
        # if height:
        row.height = Cm(height)

def write_table(table, row, col, text, text_vertical=WD_ALIGN_VERTICAL.BOTTOM, bold=False, font_size=10):
    if len(table.rows) < row:
        add_row(table, row)
    cell = table.cell(row-1, col-1)
    cell.vertical_alignment = text_vertical
    fill_cell(cell, text, bold=bold, font_size=font_size, text_alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
def insert_pic(table, row, col, pic_path, width=15, height=4):
    if len(table.rows) < row:
        add_row(table, row)
    inline_shape = table.cell(row-1, col-1).add_paragraph().add_run().add_picture(pic_path)
    if width:
        inline_shape.width = Cm(width)
    if height:
        inline_shape.height = Cm(height)
    # table.cell(row-1, col-1).add_paragraph().add_run().add_picture(pic_path, width=width)


# %%
def ifmatch(text):
    match = re.search('\{(\S+)\}',text)
    if match:
        return match.group(1)
    else:
        return None
def sub_cell(raw, new, text):
    match = re.sub('\{'+f'{raw}'+'\}', new, text)
    return match


# %%
def fill_template_table(table, data, text_alignment=None):
    out_dict = {}
    for rowIdx, row in enumerate(table.rows):
        for colIdx, cell in enumerate(table.row_cells(rowIdx)):
            match = ifmatch(cell.text)
            if match:
                if match in data.keys():
                    out_dict[match] = [rowIdx, colIdx]
                    try:
                        text = sub_cell(match, str(data[match]), cell.text)
                        if text:
                            text = str(text).split()[0]
                        cell.text = ''
                        fill_cell(cell, text, text_alignment=text_alignment)
                    except Exception as e:
                        logging.error(e)
                # else:
                #     cell.text = ''
    return out_dict
                        
def fill_template(document, data):
    tables = document.tables
    for table in tables:
        match_idx = fill_template_table(table, data)


# %%
def get_excel_info(excel_file, sheet_name='家系', index_col=['家系编号']):
    df = pd.read_excel(excel_file, sheet_name=sheet_name).fillna('').astype(str)
    df['index'] = df[index_col].apply(lambda x: ':'.join(x.to_list()), axis=1)
    return df.set_index('index')

def get_config(config_excel, sheet_name=None):
    if not sheet_name:
        df_config = pd.read_excel(config_excel, header=None, index_col=0).fillna('')[1].to_dict()
    else:
        df_config = pd.read_excel(config_excel, sheet_name, header=None, index_col=0).fillna('')[1].to_dict()
    if 'textAlignment' in df_config.keys():
        if df_config['textAlignment'].upper() == 'CENTER':
            df_config['textAlignment'] = WD_ALIGN_PARAGRAPH.CENTER
        else:
            df_config['textAlignment'] = WD_ALIGN_PARAGRAPH.LEFT
    return df_config

def fill_sample_table(table, data, text_alignment=None):
    match_idx = None
    n = 1
    for idx, sample in data.items():
        logging.info(f'\t\tsample idx:\t{idx}')
        if match_idx:
            for match, match_tableIdx in match_idx.items():
                try:
                    rowIdx, colIdx = match_tableIdx
                    add_row(table, rowIdx+n, height=1)
                    fill_cell(table.row_cells(rowIdx+n)[colIdx], sample[match], text_alignment=text_alignment)
                except Exception as e:
                    logging.error(e)
            n += 1
        else:
            match_idx = fill_template_table(table, sample, text_alignment=text_alignment)
            
def insert_sample_figure(data, table, rowIdx, colIdx, preffix, suffix, picWid, picHei, pic_path):
    n = 0
    for idx, sample in data.items():
        sample_name = idx.split(':')[1]
        for pic_preffix in preffix.split('|'):
            pic_file = os.path.join(pic_path, str(pic_preffix)+str(sample_name)+str(suffix))
            if os.path.isfile(pic_file):
                break
        # pic_file = os.path.join(pic_path, str(preffix)+str(sample_name)+str(suffix))
        if os.path.isfile(pic_file):
            insert_pic(table, rowIdx+n, colIdx, pic_file, width=picWid, height=picHei)
        n += 1
            
def make_report_2(family_info, sample_info, config, template_docx, outdir, pic_dir, png_suffix):
    config['sampleFigPreffix'] = ''
    for idx in family_info.index:
        logging.info(f'\tfamily idx:\t{idx}')
        d = Document(template_docx)
        tables = d.tables
        fill_template(d, family_info.loc[idx].to_dict())
        if idx in sample_info[config['fileName']].to_list():
            sample_info_tmp = sample_info[sample_info[config['fileName']] == idx].to_dict(orient='index')
            for i in config['sampleTable'].split('+'):
                table = tables[int(i)-1]
                fill_sample_table(table, sample_info_tmp, config['textAlignment'])
            if 'sampleFigTable' in config.keys() and config['sampleFigTable']:
                if os.path.isdir(pic_dir):
                    
                    insert_sample_figure(sample_info_tmp, tables[config['sampleFigTable']-1], config['sampleFigTableRow'], 
                                         config['sampleFigTableCol'], config['sampleFigPreffix'], png_suffix, #config['sampleFigSuffix'], 
                                         config['sampleFigwidth'], config['sampleFigheight'], pic_dir)
        d.save(os.path.join(outdir, str(idx)+'.docx'))


# %%
def excel2docx(config_file, input_file, template_docx, outdir, pic_dir=None):
    try:
        config = get_config(config_file)
        family_info = get_excel_info(input_file, sheet_name=config['fileSheetName'], index_col=[config['fileName']])
        sample_info = get_excel_info(input_file, sheet_name=config['sampleSheetName'], index_col=[config['fileName'], config['sampleName']])
        make_report_2(family_info, sample_info, config, template_docx, outdir, pic_dir)
    except Exception as e:
        logging.error(e)


# %%
if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
    bin_dir = os.path.split(os.path.realpath(__file__))[0]
    parse = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parse.add_argument('-t', '--template', default=os.path.join(bin_dir, 'template_default.docx'), help='template docx')
    parse.add_argument('-i', '--input', default=os.path.join(bin_dir, 'test_data', 'test_input.xlsx'), help='input excel')
    parse.add_argument('-c', '--config', default=os.path.join(bin_dir, 'template_config_default.xlsx'), help='config excle')
    parse.add_argument('-o', '--outdir', default=os.path.join(bin_dir, 'test_data/'), help='outdir')
    parse.add_argument('-p', '--picdir', default=os.path.join(bin_dir, 'test_data/'), help='figure dir')
    args = parse.parse_args()
    

    template_docx = args.template
    in_data = args.input
    config_data = args.config
    outdir = args.outdir
    picdir = args.picdir

    excel2docx(config_data, in_data, template_docx, outdir, picdir)

# %%

# %%
