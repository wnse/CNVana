# ---
# jupyter:
#   jupytext:
#     formats: ipynb,py:percent
#     text_representation:
#       extension: .py
#       format_name: percent
#       format_version: '1.3'
#       jupytext_version: 1.13.7
#   kernelspec:
#     display_name: Python 3 (ipykernel)
#     language: python
#     name: python3
# ---

# %%
import os
import sys
import re
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
import pandas as pd
import logging
import argparse


# %%
def fill_cell(cell, content, font_name='微软雅黑', font_size=10, color=RGBColor(0, 0, 0), bold=False):
    run = cell.paragraphs[0].add_run(str(content))
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color

def add_row(table, row_no, height=None):
    while len(table.rows) < row_no:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT.AT_LEAST
        if height:
            row.height = Cm(height)

def write_table(table, row, col, text, 
    text_vertical=WD_ALIGN_VERTICAL.CENTER, text_para=WD_ALIGN_PARAGRAPH.LEFT, 
    bold=False, font_size=10, line_spacing=16, space_after=8, space_before=8,row_height=None,
    ):
    if len(table.rows) < row:
        add_row(table, row, height=row_height)
    cell = table.cell(row-1, col-1)
    cell.vertical_alignment = text_vertical
    paragraph = cell.paragraphs[0]
    paragraph.alignment = text_para
    # paragraph.paragraph_format.line_spacing_rule = SINGL
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    # paragraph.paragraph_format.space_before = Pt(space_before)
    # paragraph.paragraph_format.space_after = Pt(space_after)

    fill_cell(cell, text, bold=bold, font_size=font_size)
    
def insert_pic(table, row, col, pic_path, width=Cm(18), height=Cm(6)):
    if len(table.rows) < row:
        add_row(table, row)
    inline_shape = table.cell(row-1, col-1).add_paragraph().add_run().add_picture(pic_path)
    inline_shape.width = width
    inline_shape.height = height


# %%
def make_report_1(tempate_file, dict_family, dict_sample, dict_config, outdir='./', png_dir=None, png_name=1):
    for f, family in dict_family.items():
        logging.info(f'{f}:{family}')
        f_out = os.path.join(outdir, f'{f}.docx')
        if f in dict_sample.keys():
            d = Document(tempate_file)
            tables = d.tables
            for idx, idx_info in family.items():
                if idx in dict_config.keys():
                    # logging.info(f'{idx}\t{idx_info}')
                    if idx_info:
                        tmp_info = f'{idx_info}'.split()[0]
                    else:
                        tmp_info = f'{idx_info}'
                    write_table(tables[dict_config[idx]['tab']-1], 
                        dict_config[idx]['row'], 
                        dict_config[idx]['col'], 
                        tmp_info,
                        text_para=WD_ALIGN_PARAGRAPH.CENTER,
                        )
            idx = '样本编号'
            row_no = 0
            for s, s_info in dict_sample[f].items():
                logging.info(f'\t{s}:{s_info}')
                s_info['样本编号'] = str(s)    
                bold = False
                if '是否推荐' in s_info.keys():
                    # if re.search(',推荐移植', s_info['备注']):
                    if s_info['是否推荐'] == '推荐移植':
                        bold = True
                for idx, idx_info in s_info.items():
                    if idx in dict_config.keys():
                        tables[dict_config[idx]['tab']-1].autofit = False
                        write_table(tables[dict_config[idx]['tab']-1], 
                            dict_config[idx]['row']+row_no, 
                            dict_config[idx]['col'], 
                            f'{idx_info}', 
                            bold=bold, 
                            text_para=WD_ALIGN_PARAGRAPH.CENTER,
                            row_height=1
                            )
                if png_dir:
                    write_table(tables[4], row_no*3+1, 1, f"样本编号：", bold=True, font_size=11)
                    write_table(tables[4], row_no*3+2, 1, f"检测结果：", bold=True, font_size=11)
                    write_table(tables[4], row_no*3+3, 1, f"染色体拷贝数结果", bold=True, font_size=11)

                    write_table(tables[4], row_no*3+1, 1, f"{s_info['样本编号']}")
                    write_table(tables[4], row_no*3+2, 1, f"{s_info['检测结果']}")
                    if png_name == 2:
                        pngName = '2'
                    elif png_name == 3:
                        pngName = '3'
                    else:
                        pngName = '1'
                    # pic_path_a = os.path.join(pngdir, f'{s}.{pngName}.png')
                    # pic_path_b = os.path.join(pngdir, f'{s}.{pngName}.png')
                    # if os.path.isfile(pic_path_a):
                    #     pic_path = pic_path_a
                    # elif os.path.isfile(pic_path_b):
                    #     pic_path = pic_path_b
                    # else:
                    #     # pic_path = None
                    pic_path = os.path.join(png_dir, f'{s}.{pngName}.png')
                    try:
                        insert_pic(tables[4], row_no*3+3, 1, pic_path)
                    except Exception as e:
                        logging.error(e)
                row_no += 1
            d.save(f_out)


# %%
# if __name__ == '__main__':
#     bin_dir = os.path.split(os.path.realpath(__file__))[0]
#     parse = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
#     parse.add_argument('-i', '--input', required=True, help='input excle required for family and sample information')
#     parse.add_argument('-c', '--config', default=os.path.join(bin_dir, 'template_config.xlsx'), help='config excel for table information')
#     parse.add_argument('-t', '--template', default=os.path.join(bin_dir, 'template.docx'), help='report template docx file')
#     parse.add_argument('-o', '--outdir', default='./', help='report output directory')
#     parse.add_argument('-p', '--png_name', default=1, choices=[1, 2, 3], type=int, help='1 for default, 2 for 2color')
#     parse.add_argument('-f', '--figure_dir', default=None, help='figure directory')
#     args = parse.parse_args()
    
#     logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
    
#     f_config = args.config
#     try:
#         df_config = pd.read_excel(f_config, index_col=0)
#         dict_config = df_config.to_dict(orient='index')
#     except Exception as e:
#         sys.exit(e)
        
#     f_input = args.input
#     try:
#         df_family = pd.read_excel(f_input, '家系').fillna('')
#         dict_family = df_family.set_index('家系编号').to_dict(orient='index')
#         # logging.info(dict_family)
#         df_sample = pd.read_excel(f_input, '样本').fillna('')
#         dict_sample = df_sample.groupby('家系编号').apply(lambda x: x.set_index('样本编号').to_dict(orient='index'))
#         # logging.info(dict_sample)
#     except Exception as e:
#         sys.exit(e)
        
#     outdir = args.outdir
#     pngdir = args.figure_dir
#     tempate_file = args.template
#     try:
#         make_PGS_report(dict_family, dict_sample, dict_config, outdir=outdir, png_dir=pngdir, png_name=args.png_name)
#     except Exception as e:
#         sys.exit(e)

# %%
